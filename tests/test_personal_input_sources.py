from __future__ import annotations

from io import BytesIO
from pathlib import Path

from fastapi.testclient import TestClient

from personal_agent.core.database import connect
from personal_agent.app import create_personal_app
from personal_agent.input_documents import MAX_UPLOAD_BYTES


def test_personal_text_source_crud_and_active_selection(tmp_path: Path, monkeypatch) -> None:
    monkeypatch.setenv("PERSONAL_AGENT_LLM_PROVIDER", "fake")
    client = _client(tmp_path)

    first = client.post(
        "/api/personal/sources/text",
        json={"title": "制动需求", "content": "# 背景\n制动信号需要诊断。"},
    )
    assert first.status_code == 200
    assert first.json()["source_type"] == "text"
    assert first.json()["title"] == "制动需求"
    assert first.json()["is_active"] is True
    assert first.json()["sections"][0]["heading"] == "背景"

    second = client.post(
        "/api/personal/sources/text",
        json={"title": "热管理需求", "content": "热管理降额需要边界条件。", "make_active": False},
    )
    assert second.status_code == 200
    assert second.json()["is_active"] is False

    sources = client.get("/api/personal/sources")
    assert sources.status_code == 200
    assert [item["title"] for item in sources.json()] == ["制动需求", "热管理需求"]

    activated = client.post(f"/api/personal/sources/{second.json()['source_uid']}/activate")
    assert activated.status_code == 200
    assert activated.json()["is_active"] is True

    refreshed = client.get("/api/personal/sources").json()
    assert refreshed[0]["source_uid"] == second.json()["source_uid"]

    detail = client.get(f"/api/personal/sources/{second.json()['source_uid']}")
    assert detail.status_code == 200
    assert "热管理降额" in detail.json()["plain_text"]

    deleted = client.delete(f"/api/personal/sources/{second.json()['source_uid']}")
    assert deleted.status_code == 200
    assert deleted.json()["status"] == "deleted"
    assert deleted.json()["active_source_uid"] == first.json()["source_uid"]

    after_delete = client.get("/api/personal/sources").json()
    assert [item["title"] for item in after_delete] == ["制动需求"]
    assert after_delete[0]["is_active"] is True
    assert client.get(f"/api/personal/sources/{second.json()['source_uid']}").status_code == 404


def test_personal_upload_sources_parse_supported_files(tmp_path: Path, monkeypatch) -> None:
    monkeypatch.setenv("PERSONAL_AGENT_LLM_PROVIDER", "fake")
    client = _client(tmp_path)

    cases = [
        ("note.txt", b"plain text requirement"),
        ("spec.md", "# 功能\n需要 DTC 场景".encode("utf-8")),
        ("input.docx", _docx_bytes("DOCX requirement", [["ID", "REQ-1"], ["Risk", "Low"]])),
        ("table.xlsx", _xlsx_bytes()),
        ("scan.pdf", _minimal_pdf_bytes("PDF requirement")),
    ]
    for filename, content in cases:
        response = client.post(
            "/api/personal/sources/upload",
            files={"file": (filename, content, "application/octet-stream")},
        )
        assert response.status_code == 200, response.text
        payload = response.json()
        assert payload["source_uid"]
        assert payload["source_type"] == Path(filename).suffix.lower().lstrip(".")
        assert payload["plain_text"].strip()
        assert payload["metadata"]["original_name"] == filename

    sources = client.get("/api/personal/sources").json()
    assert len(sources) == len(cases)
    assert sources[0]["title"] == "scan"


def test_personal_upload_rejects_doc_and_does_not_create_artifact_or_knowledge(tmp_path: Path, monkeypatch) -> None:
    monkeypatch.setenv("PERSONAL_AGENT_LLM_PROVIDER", "fake")
    db_path = tmp_path / "agent.db"
    workspace = tmp_path / "workspace"
    workspace.mkdir()
    client = TestClient(create_personal_app(db_path, workspace))

    rejected = client.post(
        "/api/personal/sources/upload",
        files={"file": ("legacy.doc", b"not supported", "application/msword")},
    )
    assert rejected.status_code == 400
    assert ".doc" in rejected.json()["detail"]
    assert ".docx" in rejected.json()["detail"]

    created = client.post(
        "/api/personal/sources/text",
        json={"title": "只作为输入", "content": "不要创建 artifact，也不要导入知识库。"},
    )
    assert created.status_code == 200

    with connect(db_path) as conn:
        source_count = conn.execute("SELECT COUNT(*) FROM personal_input_sources").fetchone()[0]
        artifact_count = conn.execute("SELECT COUNT(*) FROM artifacts").fetchone()[0]
        knowledge_count = conn.execute("SELECT COUNT(*) FROM knowledge_documents").fetchone()[0]
    assert source_count == 1
    assert artifact_count == 0
    assert knowledge_count == 0


def test_personal_upload_rejects_oversized_file(tmp_path: Path, monkeypatch) -> None:
    monkeypatch.setenv("PERSONAL_AGENT_LLM_PROVIDER", "fake")
    client = _client(tmp_path)

    response = client.post(
        "/api/personal/sources/upload",
        files={"file": ("huge.txt", b"a" * (MAX_UPLOAD_BYTES + 1), "text/plain")},
    )

    assert response.status_code == 400
    assert "too large" in response.json()["detail"]


def _client(tmp_path: Path) -> TestClient:
    db_path = tmp_path / "agent.db"
    workspace = tmp_path / "workspace"
    workspace.mkdir()
    return TestClient(create_personal_app(db_path, workspace))


def _docx_bytes(paragraph: str, table_rows: list[list[str]]) -> bytes:
    from docx import Document

    document = Document()
    document.add_paragraph(paragraph)
    table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))
    for row_index, row in enumerate(table_rows):
        for col_index, value in enumerate(row):
            table.cell(row_index, col_index).text = value
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _xlsx_bytes() -> bytes:
    from openpyxl import Workbook

    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Requirements"
    sheet.append(["ID", "Description"])
    sheet.append(["REQ-1", "xlsx requirement"])
    output = BytesIO()
    workbook.save(output)
    return output.getvalue()


def _minimal_pdf_bytes(text: str) -> bytes:
    escaped = text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
    stream = f"BT /F1 12 Tf 72 720 Td ({escaped}) Tj ET"
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >>",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
        f"<< /Length {len(stream.encode('latin-1'))} >>\nstream\n{stream}\nendstream".encode("latin-1"),
    ]
    output = BytesIO()
    output.write(b"%PDF-1.4\n")
    offsets = [0]
    for index, obj in enumerate(objects, start=1):
        offsets.append(output.tell())
        output.write(f"{index} 0 obj\n".encode("ascii"))
        output.write(obj)
        output.write(b"\nendobj\n")
    xref = output.tell()
    output.write(f"xref\n0 {len(objects) + 1}\n".encode("ascii"))
    output.write(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        output.write(f"{offset:010d} 00000 n \n".encode("ascii"))
    output.write(
        f"trailer << /Size {len(objects) + 1} /Root 1 0 R >>\nstartxref\n{xref}\n%%EOF\n".encode("ascii")
    )
    return output.getvalue()
