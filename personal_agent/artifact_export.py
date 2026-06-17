from __future__ import annotations

import json
import re
from pathlib import Path
from typing import Any

from personal_agent.core.database import connect
from personal_agent.core.utils import json_dumps, utc_now

from .artifact_drafts import get_artifact_content


FORMAT_EXTENSIONS = {
    "md": ".md",
    "docx": ".docx",
    "xlsx": ".xlsx",
    "diff": ".diff",
}


def export_personal_artifact(
    db_path: Path,
    *,
    workspace: Path,
    project_id: int,
    draft_uid: str,
    export_format: str = "",
    revision_index: int | None = None,
) -> dict[str, Any]:
    content = get_artifact_content(db_path, project_id=project_id, draft_uid=draft_uid, revision_index=revision_index)
    resolved_format = _resolve_export_format(content, export_format)
    export_dir = workspace.expanduser().resolve() / ".personal_agent" / "exports" / draft_uid
    export_dir.mkdir(parents=True, exist_ok=True)
    filename = _export_filename(content, resolved_format)
    file_path = export_dir / filename
    _write_export_file(file_path, resolved_format, content)
    exported = {
        "status": "exported",
        "draft_uid": draft_uid,
        "document_type": content["document_type"],
        "content_format": content["content_format"],
        "revision_index": content["revision"]["revision_index"],
        "export_format": resolved_format,
        "file_name": filename,
        "file_path": str(file_path),
        "download_url": f"/api/personal/drafts/{draft_uid}/download?format={resolved_format}",
        "boundaries": {
            "personal_export_only": True,
            "writes_release_record": False,
            "applies_code": False,
        },
    }
    _record_export_metadata(db_path, project_id=project_id, export=exported)
    return exported


def resolve_personal_artifact_download(
    db_path: Path,
    *,
    workspace: Path,
    project_id: int,
    draft_uid: str,
    export_format: str = "",
) -> dict[str, Any]:
    content = get_artifact_content(db_path, project_id=project_id, draft_uid=draft_uid)
    resolved_format = _resolve_export_format(content, export_format)
    file_path = workspace.expanduser().resolve() / ".personal_agent" / "exports" / draft_uid / _export_filename(content, resolved_format)
    if not file_path.exists():
        export_personal_artifact(
            db_path,
            workspace=workspace,
            project_id=project_id,
            draft_uid=draft_uid,
            export_format=resolved_format,
        )
    media_type = {
        "md": "text/markdown; charset=utf-8",
        "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        "diff": "text/x-diff; charset=utf-8",
    }[resolved_format]
    return {"file_path": str(file_path), "file_name": file_path.name, "media_type": media_type, "export_format": resolved_format}


def _resolve_export_format(content: dict[str, Any], requested: str) -> str:
    requested = requested.strip().lower().lstrip(".")
    allowed = _allowed_formats(content)
    if not requested:
        return allowed[0]
    if requested not in allowed:
        raise ValueError(f"export format {requested} is not supported for {content['document_type']} ({content['content_format']})")
    return requested


def _allowed_formats(content: dict[str, Any]) -> list[str]:
    document_type = str(content.get("document_type") or "")
    content_format = str(content.get("content_format") or "")
    if content_format == "diff" or document_type in {"c_code_diff", "unit_test_code_or_diff"}:
        return ["diff"]
    if content_format == "json_table" or document_type == "test_case_spec":
        return ["xlsx"]
    if content_format in {"markdown", "text"}:
        return ["md", "docx"]
    raise ValueError(f"unsupported content_format for export: {content_format}")


def _export_filename(content: dict[str, Any], export_format: str) -> str:
    title = _safe_filename(str(content.get("title") or content.get("document_type") or "draft"))
    revision = int((content.get("revision") or {}).get("revision_index") or content.get("current_revision") or 1)
    return f"{title}_v{revision}{FORMAT_EXTENSIONS[export_format]}"


def _write_export_file(file_path: Path, export_format: str, content: dict[str, Any]) -> None:
    text = str(content.get("content") or "")
    if export_format in {"md", "diff"}:
        file_path.write_text(text, encoding="utf-8")
        return
    if export_format == "docx":
        _write_docx(file_path, text)
        return
    if export_format == "xlsx":
        _write_xlsx(file_path, text)
        return
    raise ValueError(f"unsupported export format: {export_format}")


def _write_docx(file_path: Path, markdown: str) -> None:
    try:
        from docx import Document
    except Exception as exc:  # pragma: no cover - dependency is installed in normal app env
        raise ValueError("python-docx is required to export .docx") from exc
    document = Document()
    for raw_line in markdown.splitlines():
        line = raw_line.rstrip()
        if not line:
            continue
        if line.startswith("# "):
            document.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            document.add_heading(line[3:].strip(), level=2)
        elif line.startswith("### "):
            document.add_heading(line[4:].strip(), level=3)
        elif line.startswith("- "):
            document.add_paragraph(line[2:].strip(), style="List Bullet")
        elif re.match(r"^\d+\.\s+", line):
            document.add_paragraph(re.sub(r"^\d+\.\s+", "", line).strip(), style="List Number")
        else:
            document.add_paragraph(line)
    document.save(file_path)


def _write_xlsx(file_path: Path, json_table: str) -> None:
    try:
        from openpyxl import Workbook
    except Exception as exc:  # pragma: no cover - dependency is installed in normal app env
        raise ValueError("openpyxl is required to export .xlsx") from exc
    payload = json.loads(json_table or "{}")
    columns = payload.get("columns") or []
    rows = payload.get("rows") or []
    if not isinstance(columns, list) or not isinstance(rows, list):
        raise ValueError("json_table export requires columns and rows")
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "draft"
    for col_index, column in enumerate(columns, start=1):
        sheet.cell(row=1, column=col_index, value=str(column))
    for row_index, row in enumerate(rows, start=2):
        if isinstance(row, dict):
            values = [row.get(str(column), "") for column in columns]
        elif isinstance(row, list):
            values = row
        else:
            values = [str(row)]
        for col_index, value in enumerate(values, start=1):
            sheet.cell(row=row_index, column=col_index, value=json_dumps(value) if isinstance(value, (dict, list)) else value)
    workbook.save(file_path)


def _record_export_metadata(db_path: Path, *, project_id: int, export: dict[str, Any]) -> None:
    now = utc_now()
    with connect(db_path) as conn:
        conn.execute(
            """
            INSERT INTO audit_events(project_id, event_type, message, payload_json, created_at)
            VALUES (?, 'PERSONAL_DRAFT_EXPORTED', ?, ?, ?)
            """,
            (
                project_id,
                f"导出个人草稿 {export['draft_uid']} 为 {export['export_format']}",
                json_dumps(export),
                now,
            ),
        )


def _safe_filename(value: str) -> str:
    text = re.sub(r'[<>:"/\\|?*\x00-\x1f]', "_", value).strip(" ._")
    return text[:80] or "draft"
