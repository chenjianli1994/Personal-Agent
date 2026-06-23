from __future__ import annotations

import json
from dataclasses import dataclass
from io import BytesIO
from pathlib import Path
from typing import Any
from uuid import uuid4

from personal_agent.core.database import connect
from personal_agent.core.utils import json_dumps, utc_now


SUPPORTED_SOURCE_TYPES = {"text", "txt", "md", "docx", "pdf", "xlsx", "xlsm"}
MAX_UPLOAD_BYTES = 20 * 1024 * 1024


@dataclass(frozen=True)
class ParsedInputSource:
    source_type: str
    title: str
    plain_text: str
    tables: list[dict[str, Any]]
    sections: list[dict[str, Any]]
    metadata: dict[str, Any]

    def to_payload(self, source_uid: str = "") -> dict[str, Any]:
        return {
            "source_uid": source_uid,
            "source_type": self.source_type,
            "title": self.title,
            "plain_text": self.plain_text,
            "tables": self.tables,
            "sections": self.sections,
            "metadata": self.metadata,
        }


def parse_text_source(content: str, *, title: str = "文本输入") -> ParsedInputSource:
    text = content.strip()
    if not text:
        raise ValueError("text content is required")
    return ParsedInputSource(
        source_type="text",
        title=title.strip() or "文本输入",
        plain_text=text,
        tables=[],
        sections=_sections_from_text(text),
        metadata={"original_name": "", "page_refs": [], "sheet_refs": []},
    )


def parse_uploaded_source(filename: str, content: bytes) -> ParsedInputSource:
    original_name = Path(filename or "").name
    suffix = Path(original_name).suffix.lower().lstrip(".")
    if suffix == "doc":
        raise ValueError("unsupported file type .doc; please save as .docx")
    if suffix not in {"txt", "md", "docx", "pdf", "xlsx", "xlsm"}:
        raise ValueError(f"unsupported file type: .{suffix or 'unknown'}")
    if not content:
        raise ValueError("uploaded file is empty")
    if len(content) > MAX_UPLOAD_BYTES:
        raise ValueError("uploaded file too large")
    title = Path(original_name).stem or original_name or "上传材料"
    if suffix in {"txt", "md"}:
        text = _decode_text(content).strip()
        if not text:
            raise ValueError("uploaded text file has no readable content")
        return ParsedInputSource(
            source_type=suffix,
            title=title,
            plain_text=text,
            tables=[],
            sections=_sections_from_text(text),
            metadata={"original_name": original_name, "page_refs": [], "sheet_refs": []},
        )
    if suffix == "docx":
        return _parse_docx(content, title=title, original_name=original_name)
    if suffix == "pdf":
        return _parse_pdf(content, title=title, original_name=original_name)
    return _parse_xlsx(content, title=title, original_name=original_name, source_type=suffix)


def create_input_source(
    db_path: Path,
    *,
    project_id: int,
    parsed: ParsedInputSource,
    make_active: bool = True,
) -> dict[str, Any]:
    source_uid = f"src_{uuid4().hex}"
    now = utc_now()
    with connect(db_path) as conn:
        if make_active:
            conn.execute(
                "UPDATE personal_input_sources SET is_active=0 WHERE project_id=?",
                (project_id,),
            )
        conn.execute(
            """
            INSERT INTO personal_input_sources(
                source_uid, project_id, source_type, title, plain_text,
                sections_json, tables_json, metadata_json, status, is_active,
                created_at, updated_at
            )
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, 'active', ?, ?, ?)
            """,
            (
                source_uid,
                project_id,
                parsed.source_type,
                parsed.title,
                parsed.plain_text,
                json_dumps(parsed.sections),
                json_dumps(parsed.tables),
                json_dumps(parsed.metadata),
                1 if make_active else 0,
                now,
                now,
            ),
        )
    return get_input_source(db_path, project_id=project_id, source_uid=source_uid)


def list_input_sources(db_path: Path, *, project_id: int) -> list[dict[str, Any]]:
    with connect(db_path) as conn:
        rows = conn.execute(
            """
            SELECT * FROM personal_input_sources
            WHERE project_id=? AND status='active'
            ORDER BY is_active DESC, id DESC
            """,
            (project_id,),
        ).fetchall()
    return [_row_to_source(row) for row in rows]


def get_input_source(db_path: Path, *, project_id: int, source_uid: str) -> dict[str, Any]:
    with connect(db_path) as conn:
        row = conn.execute(
            """
            SELECT * FROM personal_input_sources
            WHERE project_id=? AND source_uid=? AND status='active'
            """,
            (project_id, source_uid),
        ).fetchone()
    if row is None:
        raise ValueError("source not found")
    return _row_to_source(row)


def activate_input_source(db_path: Path, *, project_id: int, source_uid: str) -> dict[str, Any]:
    with connect(db_path) as conn:
        row = conn.execute(
            "SELECT id FROM personal_input_sources WHERE project_id=? AND source_uid=? AND status='active'",
            (project_id, source_uid),
        ).fetchone()
        if row is None:
            raise ValueError("source not found")
        conn.execute("UPDATE personal_input_sources SET is_active=0 WHERE project_id=?", (project_id,))
        conn.execute(
            "UPDATE personal_input_sources SET is_active=1, updated_at=? WHERE id=?",
            (utc_now(), row["id"]),
        )
    return get_input_source(db_path, project_id=project_id, source_uid=source_uid)


def activate_input_sources(db_path: Path, *, project_id: int, source_uids: list[str]) -> list[dict[str, Any]]:
    unique_uids = list(dict.fromkeys(uid.strip() for uid in source_uids if uid.strip()))
    if not unique_uids:
        return []
    if len(unique_uids) > 5:
        raise ValueError("at most 5 input files can be attached to one message")
    placeholders = ",".join("?" for _ in unique_uids)
    now = utc_now()
    with connect(db_path) as conn:
        rows = conn.execute(
            f"""
            SELECT id, source_uid FROM personal_input_sources
            WHERE project_id=? AND status='active' AND source_uid IN ({placeholders})
            """,
            (project_id, *unique_uids),
        ).fetchall()
        found = {str(row["source_uid"]): int(row["id"]) for row in rows}
        missing = [uid for uid in unique_uids if uid not in found]
        if missing:
            raise ValueError(f"source not found: {missing[0]}")
        conn.execute("UPDATE personal_input_sources SET is_active=0 WHERE project_id=?", (project_id,))
        conn.executemany(
            "UPDATE personal_input_sources SET is_active=1, updated_at=? WHERE id=?",
            [(now, found[uid]) for uid in unique_uids],
        )
    return [get_input_source(db_path, project_id=project_id, source_uid=uid) for uid in unique_uids]


def delete_input_source(db_path: Path, *, project_id: int, source_uid: str) -> dict[str, Any]:
    now = utc_now()
    with connect(db_path) as conn:
        row = conn.execute(
            "SELECT id, is_active FROM personal_input_sources WHERE project_id=? AND source_uid=? AND status='active'",
            (project_id, source_uid),
        ).fetchone()
        if row is None:
            raise ValueError("source not found")
        conn.execute(
            "UPDATE personal_input_sources SET status='deleted', is_active=0, updated_at=? WHERE id=?",
            (now, row["id"]),
        )
        active_uid = ""
        if bool(row["is_active"]):
            fallback = conn.execute(
                """
                SELECT id, source_uid FROM personal_input_sources
                WHERE project_id=? AND status='active'
                ORDER BY id DESC LIMIT 1
                """,
                (project_id,),
            ).fetchone()
            if fallback is not None:
                active_uid = str(fallback["source_uid"])
                conn.execute(
                    "UPDATE personal_input_sources SET is_active=1, updated_at=? WHERE id=?",
                    (now, fallback["id"]),
                )
    return {"status": "deleted", "source_uid": source_uid, "active_source_uid": active_uid}


def _parse_docx(content: bytes, *, title: str, original_name: str) -> ParsedInputSource:
    try:
        from docx import Document
    except ImportError as exc:
        raise ValueError("python-docx is required to parse .docx files") from exc
    document = Document(BytesIO(content))
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    tables: list[dict[str, Any]] = []
    for table_index, table in enumerate(document.tables):
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        if rows:
            tables.append({"table_index": table_index, "rows": rows})
    table_text = _tables_to_text(tables)
    text = "\n".join(paragraphs + ([table_text] if table_text else [])).strip()
    if not text:
        raise ValueError("docx has no readable text")
    return ParsedInputSource(
        source_type="docx",
        title=title,
        plain_text=text,
        tables=tables,
        sections=_sections_from_text(text),
        metadata={"original_name": original_name, "page_refs": [], "sheet_refs": []},
    )


def _parse_pdf(content: bytes, *, title: str, original_name: str) -> ParsedInputSource:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise ValueError("pypdf is required to parse .pdf files") from exc
    reader = PdfReader(BytesIO(content))
    page_refs: list[dict[str, Any]] = []
    pages: list[str] = []
    for index, page in enumerate(reader.pages, start=1):
        page_text = (page.extract_text() or "").strip()
        if page_text:
            page_refs.append({"page": index})
            pages.append(page_text)
    text = "\n\n".join(pages).strip()
    if not text:
        raise ValueError("pdf has no extractable text")
    return ParsedInputSource(
        source_type="pdf",
        title=title,
        plain_text=text,
        tables=[],
        sections=_sections_from_text(text),
        metadata={"original_name": original_name, "page_refs": page_refs, "sheet_refs": []},
    )


def _parse_xlsx(content: bytes, *, title: str, original_name: str, source_type: str) -> ParsedInputSource:
    try:
        from openpyxl import load_workbook
    except ImportError as exc:
        raise ValueError("openpyxl is required to parse .xlsx/.xlsm files") from exc
    workbook = load_workbook(BytesIO(content), data_only=True, read_only=True)
    tables: list[dict[str, Any]] = []
    sheet_refs: list[dict[str, Any]] = []
    text_parts: list[str] = []
    for sheet in workbook.worksheets:
        rows: list[list[str]] = []
        for row in sheet.iter_rows(values_only=True):
            values = ["" if value is None else str(value) for value in row]
            if any(cell.strip() for cell in values):
                rows.append(values)
        if not rows:
            continue
        sheet_refs.append({"sheet": sheet.title, "row_count": len(rows)})
        tables.append({"sheet_name": sheet.title, "rows": rows})
        text_parts.append(f"# {sheet.title}\n" + "\n".join("\t".join(row) for row in rows))
    text = "\n\n".join(text_parts).strip()
    if not text:
        raise ValueError("spreadsheet has no readable cells")
    return ParsedInputSource(
        source_type=source_type,
        title=title,
        plain_text=text,
        tables=tables,
        sections=_sections_from_text(text),
        metadata={"original_name": original_name, "page_refs": [], "sheet_refs": sheet_refs},
    )


def _decode_text(content: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "gb18030"):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    return content.decode("utf-8", errors="replace")


def _sections_from_text(text: str) -> list[dict[str, Any]]:
    sections: list[dict[str, Any]] = []
    current_heading = "正文"
    current_lines: list[str] = []
    for line in text.splitlines():
        stripped = line.strip()
        if stripped.startswith("#"):
            if current_lines:
                sections.append({"heading": current_heading, "content": "\n".join(current_lines).strip()})
                current_lines = []
            current_heading = stripped.lstrip("#").strip() or "未命名章节"
        else:
            current_lines.append(line)
    if current_lines:
        sections.append({"heading": current_heading, "content": "\n".join(current_lines).strip()})
    if not sections and text.strip():
        sections.append({"heading": "正文", "content": text.strip()})
    return [item for item in sections if item["content"]]


def _tables_to_text(tables: list[dict[str, Any]]) -> str:
    parts: list[str] = []
    for table in tables:
        rows = table.get("rows") or []
        if rows:
            parts.append("\n".join("\t".join(str(cell) for cell in row) for row in rows))
    return "\n\n".join(parts)


def _row_to_source(row: Any) -> dict[str, Any]:
    payload = {
        "id": row["id"],
        "source_uid": row["source_uid"],
        "project_id": row["project_id"],
        "source_type": row["source_type"],
        "title": row["title"],
        "plain_text": row["plain_text"],
        "tables": _loads_json(row["tables_json"], []),
        "sections": _loads_json(row["sections_json"], []),
        "metadata": _loads_json(row["metadata_json"], {}),
        "status": row["status"],
        "is_active": bool(row["is_active"]),
        "created_at": row["created_at"],
        "updated_at": row["updated_at"],
    }
    payload["preview"] = payload["plain_text"][:500]
    return payload


def _loads_json(text: str, default: Any) -> Any:
    try:
        return json.loads(text or "")
    except Exception:
        return default
